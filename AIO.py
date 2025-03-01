import streamlit as st
import asyncio
import aiohttp
from typing import List, Dict, Any, Optional
from pydantic import BaseModel, ValidationError, Field
from crewai import Crew, Agent, Task
from langchain_groq import ChatGroq
from langchain_google_genai import ChatGoogleGenerativeAI
import warnings
import logging
import json
import os
from textwrap import dedent

# Suppress specific warnings
warnings.filterwarnings("ignore", category=UserWarning, module="pydantic")
warnings.filterwarnings("ignore", category=SyntaxWarning)

# Configure logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

# --- Pydantic Models ---
class ActionItem(BaseModel):
    description: str = Field(..., description="Description of the action item.")
    priority: str = Field(..., description="Priority level of the action (e.g., High, Medium, Low).")
    compliance_status: str = Field(..., description="Status of compliance (e.g., Compliant, Non-compliant, Partially compliant).")
    timeline: str = Field(..., description="Timeline for the action item.")

class RiskMitigation(BaseModel):
    strategy: str = Field(..., description="The mitigation strategy.")
    stakeholders: List[str] = Field(..., description="List of stakeholders involved.")
    steps: List[str] = Field(..., description="Steps for the mitigation strategy.")

class ParagraphAnalysis(BaseModel):
    text: str = Field(..., description="The text of the paragraph.")
    actions: List[ActionItem] = Field(..., description="List of action items extracted from the paragraph.")
    mitigations: List[RiskMitigation] = Field(default_factory=list, description="List of risk mitigations.")  # Optional mitigations

class DocumentAnalysis(BaseModel):
    context: Dict[str, Any] = Field(..., description="Context of the document.")
    paragraphs: List[ParagraphAnalysis] = Field(..., description="Analysis of each paragraph.")
    report: str = Field(..., description="Final report summarizing the analysis.")

# --- LLM Factory ---
class LLMFactory:
    _instances: Dict[tuple, Any] = {}

    @classmethod
    def get_llm(cls, api_choice: str, temperature: float, model_name: str = None):
        key = (api_choice, temperature, model_name)
        if key not in cls._instances:
            if api_choice == "groq":
                model = model_name or "mixtral-8x7b-32768"  # Default model
                cls._instances[key] = ChatGroq(
                    temperature=temperature,
                    model=model,
                    groq_api_key=os.getenv("GROQ_API_KEY")
                )
            elif api_choice == "google_ai":
                model = model_name or "gemini-pro"  # Default model for Google AI
                cls._instances[key] = ChatGoogleGenerativeAI(
                    model=model,
                    temperature=temperature,
                    google_api_key=os.getenv("GOOGLE_API_KEY"),
                    convert_system_message_to_human=True
                )
            else:
                raise ValueError(f"Invalid API choice: {api_choice}")
        return cls._instances[key]


# --- Optimized Base Agent ---
class BaseAgent(Agent):
    def __init__(self, api_choice, role: str, goal: str, backstory: str, temperature:float, model_name:Optional[str] = None):
        super().__init__(
            role=role,
            goal=goal,
            backstory=backstory,
            llm=LLMFactory.get_llm(api_choice, temperature, model_name),  # Use LLMFactory
            verbose=True,
            allow_delegation=False,
            max_iter=15,  #increased for complex tasks
            max_round_robin=3,
            cache=True  # Enable caching for repeated queries
        )

# --- Agent Factory ---
class AgentFactory:
    @staticmethod
    def create_parser(api_choice: str, temperature: float, model_name: str = None):
        return BaseAgent(
            api_choice,
            "Regulatory Parser",
            "Parse regulatory documents efficiently",
            "Specialized in document structure analysis",
            temperature,
            model_name
        )

    @staticmethod
    def create_context(api_choice: str, temperature: float, industry: str, model_name: str = None):
        return BaseAgent(
            api_choice,
            "Context Analyzer",
            f"Analyze regulatory context for {industry}",
            "Expert in industry compliance",
            temperature,
            model_name
        )

    @staticmethod
    def create_action(api_choice: str, temperature: float, industry: str, model_name: str = None):
        return BaseAgent(
           api_choice,
            "Action Extractor",
            f"Extract actionable items for {industry}",
            "Compliance action specialist",
            temperature,
            model_name

        )

    @staticmethod
    def create_compliance(api_choice: str, temperature: float, industry: str, model_name: str = None):
        return BaseAgent(
            api_choice,
            "Compliance Checker",
            f"Verify compliance for {industry}",
            "Regulatory compliance expert",
            temperature,
            model_name
        )

    @staticmethod
    def create_priority(api_choice: str, temperature: float, model_name: str = None):
        return BaseAgent(
           api_choice,
            "Priority Assigner",
            "Assign priority levels",
            "Risk assessment specialist",
            temperature,
            model_name
        )

    @staticmethod
    def create_mitigation(api_choice: str, temperature: float, industry: str, model_name: str = None):
        return BaseAgent(
            api_choice,
            "Risk Mitigator",
            f"Develop mitigation strategies for {industry}",
            "Risk management expert",
            temperature,
            model_name
        )

    @staticmethod
    def create_timeline(api_choice: str, temperature: float, model_name: str = None):
        return BaseAgent(
            api_choice,
            "Timeline Planner",
            "Plan implementation timelines",
            "Project scheduling expert",
            temperature,
            model_name
        )

    @staticmethod
    def create_report(api_choice: str, temperature: float, model_name: str = None):
        return BaseAgent(
            api_choice,
            "Report Generator",
            "Generate structured reports",
            "Documentation specialist",
            temperature,
            model_name
        )

# --- Core Processing Functions ---
async def parse_document(agent: BaseAgent, content: str) -> List[str]:
    logger.debug("Parsing document content")
    paragraphs =  [p.strip() for p in content.split('\n\n') if p.strip()]
    return paragraphs


async def process_paragraph(paragraph: str, agents: Dict[str, BaseAgent], context_summary: str) -> ParagraphAnalysis:
    logger.debug(f"Processing paragraph: {paragraph[:50]}...")

    try:
        action_task = Task(description=dedent(f"""
            Analyze this paragraph within the context: {context_summary}
            Paragraph content: {paragraph}
            Identify and extract specific actionable items required for compliance.
            Focus on actions that are clear, concise, and directly actionable, avoid general statements.
            """), agent=agents["action"])

        compliance_task = Task(description=dedent(f"""
            Given the context: {context_summary}
            Paragraph content: {paragraph}
            Assess the paragraph for compliance with relevant regulations.
            Identify specific areas of compliance and non-compliance.  State if it's compliant or not.
            """), agent=agents["compliance"])

        priority_task = Task(description=dedent(f"""
            Based on the context: {context_summary} and the content of the paragraph: '{paragraph}',
            assign a priority level (High, Medium, Low) to each identified action item.
            Consider the potential impact and urgency.
            """), agent=agents["priority"])

        mitigation_task = Task(description=dedent(f"""
            Analyze the paragraph within the overall context: {context_summary}
            Paragraph: {paragraph}
            Identify potential risks and develop specific, actionable mitigation strategies.  Include Stakeholders.
            """), agent=agents["mitigation"])

        timeline_task = Task(description=dedent(f"""
            For the paragraph: '{paragraph}', within the context: {context_summary}
            Develop realistic timelines for each actionable item.
            Consider dependencies and potential constraints. Be as specific as possible
            """), agent=agents["timeline"])


        crew = Crew(agents=[agents["action"], agents["compliance"], agents["priority"], agents["mitigation"], agents["timeline"]],
                    tasks=[action_task, compliance_task, priority_task, mitigation_task, timeline_task],
                    process=Crew.Process.sequential, # Sequential is better for complex inter-dependencies.
                    verbose=2)
        results = crew.kickoff()

        # Result Processing with error handling and type checking
        actions = []
        mitigations = []

        # Helper function to extract and validate
        def extract_value(results, keyword, default):
            for result in results:
                if keyword in result:
                    try:
                        # Attempt to parse as JSON if it's a string
                        if isinstance(result, str):
                            return json.loads(result)
                        return result # Return directly if it's already a dict

                    except (json.JSONDecodeError, TypeError):
                        logger.warning(f"Could not parse {keyword}: {result}")

            return default
        
        # Collect results from all tasks
        task_results = {}
        for task in [action_task, compliance_task, priority_task, mitigation_task, timeline_task]:
            for result in results:
                if result.startswith(task.description[:20]):  # Using part of the description
                    task_results[task] = result
                    break
    
        for i in range(len(results)): # Use the number of results returned
            try:
                # Extract results safely with defaults
                action_description = extract_value([task_results.get(action_task, "")], "action", {}).get("description", "No action identified")
                priority = extract_value([task_results.get(priority_task, "")], "priority", {}).get("priority", "Medium")
                compliance_status = extract_value([task_results.get(compliance_task, "")], "compliance_status", {}).get("status","Pending")
                timeline = extract_value([task_results.get(timeline_task, "")], "timeline", {}).get("timeline", "TBD")
                
                actions.append(ActionItem(description=action_description, priority=priority, compliance_status=compliance_status, timeline=timeline))
            except Exception as e:
                logger.error(f"Error creating ActionItem {i}: {e}")
                actions.append(ActionItem(description="Error processing action", priority="Medium", compliance_status="Error", timeline="Error"))
                continue # Continue to avoid crashing the whole process.

        # Mitigation Extraction (similar approach)
        mitigation_result = extract_value([task_results.get(mitigation_task, "")], "mitigation", [])
        if isinstance(mitigation_result, list):
            for m in mitigation_result:
                if isinstance(m, dict):
                    try:
                        mitigations.append(RiskMitigation(**m))
                    except ValidationError as e:
                        logger.error(f"Error processing mitigation: {m}. Error: {e}")

        return ParagraphAnalysis(text=paragraph, actions=actions, mitigations=mitigations)

    except Exception as e:
        logger.exception(f"Error processing paragraph: {paragraph[:100]}...")
        return ParagraphAnalysis(text=paragraph, actions=[], mitigations=[])  # Return a default object


# --- Main Processing Function ---
async def process_regulatory_obligation(
    input_type: str,
    input_source: str,
    api_actions: str = "groq",
    api_mitigation: str = "groq",
    temperature: float = 0.5,
    industry: str = "General",
    groq_model_name: str = None,
    google_model_name: str = None
) -> DocumentAnalysis:
    try:
        logger.debug(f"Starting processing with input type: {input_type}, source: {input_source}")

        # Load content
        if input_type == "url":
            async with aiohttp.ClientSession() as session:
                try:
                    async with session.get(input_source, timeout=aiohttp.ClientTimeout(total=10)) as resp:
                        resp.raise_for_status()  # Raise HTTPError for bad responses (4xx or 5xx)
                        document_content = await resp.text()
                except (aiohttp.ClientError, asyncio.TimeoutError) as e:
                    raise RuntimeError(f"Failed to fetch URL: {e}")
        elif input_type == "file upload":
            document_content = input_source
        else:
            raise ValueError(f"Invalid input type: {input_type}")

        if not document_content.strip():
            raise ValueError("Empty document content")

        # Initialize agents
        parser_agent = AgentFactory.create_parser(api_actions, temperature, groq_model_name or google_model_name)
        context_agent = AgentFactory.create_context(api_actions, temperature, industry, groq_model_name or google_model_name)
        agents = {
            "action": AgentFactory.create_action(api_actions, temperature, industry, groq_model_name or google_model_name),
            "compliance": AgentFactory.create_compliance(api_actions, temperature, industry, groq_model_name or google_model_name),
            "priority": AgentFactory.create_priority(api_actions, temperature, groq_model_name or google_model_name),
            "mitigation": AgentFactory.create_mitigation(api_mitigation, temperature, industry, groq_model_name or google_model_name),  # Use specified API
            "timeline": AgentFactory.create_timeline(api_actions, temperature, groq_model_name or google_model_name),
            "report": AgentFactory.create_report(api_actions, temperature, groq_model_name or google_model_name)
        }

        # Process document
        paragraphs = await parse_document(parser_agent, document_content)
        if not paragraphs:
            raise ValueError("No paragraphs found in the document.")

        # Analyze context
        context_task = Task(
            description=dedent(f"""
                Analyze the entire document and provide a concise summary of its overall context.
                Document content: {document_content}
                Identify the main regulatory themes, the scope of the regulations, and the key entities affected.
            """),
            agent=context_agent
        )

        context_crew = Crew(
            agents=[context_agent],
            tasks=[context_task],
            verbose=2
        )

        context_result = context_crew.kickoff()
        context_summary = extract_value([context_result], "summary", {}).get("Overall Context", "Context analysis failed.")
        if isinstance(context_summary, list):
            context_summary = " ".join(context_summary) # Join to get string

        # Parallel paragraph processing using asyncio.gather
        paragraph_results = await asyncio.gather(
            *(process_paragraph(para, agents, context_summary) for para in paragraphs)
        )

        # Generate report using CrewAI for consistency
        report_task = Task(
            description=dedent(f"""
                Generate a comprehensive report summarizing the analysis of the regulatory document.
                Overall context: {context_summary}
                Paragraph summaries and their extracted actions and mitigations:
                {paragraph_results}
                The report should clearly present the context, identified actions (with priorities,
                compliance status, and timelines), and proposed mitigation strategies. Structure the report for clarity and readability.
            """),
            agent=agents["report"]
        )

        report_crew = Crew(
            agents=[agents["report"]],
            tasks=[report_task],
            verbose=2
        )
        report_result_list = report_crew.kickoff()
        report_result = extract_value([report_result_list], "report", {}).get("Report", "Report generation failed.")
        if isinstance(report_result, list):
            report_result = " ".join(report_result)


        logger.debug("Document processed successfully")
        return DocumentAnalysis(
            context={"summary": context_summary},
            paragraphs=paragraph_results,
            report=report_result
        )

    except Exception as e:
        logger.exception(f"Processing failed: {str(e)}")
        raise RuntimeError(f"Processing failed: {str(e)}")

# --- Streamlit App ---
async def main_streamlit():
    st.title("Regulatory Obligation Processor")

    input_type = st.selectbox("Input Type", ["file upload", "url"])
    input_source = ""

    if input_type == "file upload":
        uploaded_file = st.file_uploader("Upload a file", type=["txt", "pdf", "docx"])  # Added file types
        if uploaded_file is not None:
            try:
                # Handle different file types
                if uploaded_file.name.endswith('.txt'):
                    input_source = uploaded_file.read().decode("utf-8")
                elif uploaded_file.name.endswith('.pdf'):
                    # Using pypdfium2 for PDF handling (more reliable)
                    import pypdfium2 as pdfium
                    pdf = pdfium.PdfDocument(uploaded_file)
                    text_pages = []
                    for page_number in range(len(pdf)):
                        page = pdf.get_page(page_number)
                        text_page = page.get_textpage()
                        text_pages.append(text_page.get_text_range())
                    input_source = "\n".join(text_pages)

                elif uploaded_file.name.endswith('.docx'):
                    from docx import Document
                    document = Document(uploaded_file)
                    input_source = "\n".join([paragraph.text for paragraph in document.paragraphs])
                else:
                    st.error("Unsupported file type.")
                    return # Early return on error
            except Exception as e:
                st.error(f"Error reading file: {e}")
                return # Exit if reading file fails

    elif input_type == "url":
        input_source = st.text_input("Enter URL")

    api_actions = st.selectbox("API for Actions", ["groq", "google_ai"])
    api_mitigation = st.selectbox("API for Mitigation", ["groq", "google_ai"])
    temperature = st.slider("Temperature", min_value=0.0, max_value=1.0, value=0.5)
    industry = st.text_input("Industry", "General")
    groq_model_name = st.selectbox("Groq Model", ["mixtral-8x7b-32768", "llama2-70b-4096"], index=0) if api_actions == 'groq' or api_mitigation == 'groq' else None
    google_model_name = st.selectbox("Google AI Model", ["gemini-pro", "gemini-1.5-pro-002"],index=0) if api_actions == 'google_ai' or api_mitigation == 'google_ai' else None


    if st.button("Process"):
        if not input_source:
            st.error("Please provide input source.")
            return # Ensure we don't proceed without input

        try:
            # Await the coroutine!
            with st.spinner("Processing..."):
                result = await process_regulatory_obligation(
                    input_type=in
